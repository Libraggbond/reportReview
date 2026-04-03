#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Unified method to extract content from a single DOCX file
Usage: python extract_docx.py <docx_file_path>
"""
from docx import Document
import sys

def extract_docx_content(docx_path, output_file=None):
    """
    Extract paragraphs and tables from a DOCX file.
    
    Args:
        docx_path: Path to the .docx file
        output_file: Optional path to save extracted content. If None, prints to stdout.
    
    Returns:
        Extracted content as a string
    """
    doc = Document(docx_path)
    lines = []
    
    # Extract document name
    import os
    doc_name = os.path.basename(docx_path)
    lines.append(f"{'='*80}")
    lines.append(f"文档: {doc_name}")
    lines.append(f"{'='*80}\n")
    
    # Extract paragraphs in batches
    lines.append("【段落内容】")
    paras = [p for p in doc.paragraphs if p.text.strip()]
    total_paras = len(paras)
    lines.append(f"总段落数: {total_paras}\n")
    
    for i, para in enumerate(paras):
        lines.append(f"[P{i}] {para.text.strip()}")
    
    lines.append("")
    
    # Extract tables
    lines.append("【表格内容】")
    total_tables = len(doc.tables)
    lines.append(f"总表格数: {total_tables}\n")
    
    for ti, table in enumerate(doc.tables):
        lines.append(f"--- 表格 {ti+1} ({len(table.rows)}行 x {len(table.columns)}列) ---")
        
        # Process table in batches of 30 rows
        rows = list(table.rows)
        for start in range(0, len(rows), 30):
            batch = rows[start:start+30]
            for ri, row in enumerate(batch, start):
                cells = [cell.text.strip().replace('\n', '\\n') for cell in row.cells]
                lines.append(f"  行{ri}: " + ' | '.join(cells))
        
        lines.append("")  # Empty line between tables
    
    content = '\n'.join(lines)
    
    if output_file:
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(content)
        print(f"内容已提取到: {output_file}")
    else:
        print(content)
    
    return content

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python extract_docx.py <docx_file_path> [output_file]")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    extract_docx_content(docx_path, output_file)