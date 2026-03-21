#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档脱敏工具

功能:
1. 删除 docx 文件中的所有图片
2. 删除所有页眉页脚
3. 文本替换:
   - 替换'青岛'、'山东'、'中国'为'乌城'
   - 替换备案编号格式为 00-00
   - 替换其他敏感信息
"""

import os
import re
import argparse
from docx import Document


def remove_images_from_docx(doc_path):
    """删除 docx 文件中的所有图片"""
    print(f"  - 删除图片：{doc_path}")

    doc = Document(doc_path)
    doc_element = doc._element

    # 查找并删除所有 blip 元素 (图片引用)
    blip_elements = doc_element.findall(
        './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip',
    )
    for blip in blip_elements:
        parent = blip.getparent()
        if parent is not None:
            parent.remove(blip)

    # 删除 drawing 元素 (包含图片的容器)
    drawing_elements = doc_element.findall(
        './/{http://schemas.openxmlformats.org/drawingml/2006/main}drawing',
    )
    for drawing in drawing_elements:
        parent = drawing.getparent()
        if parent is not None:
            parent.remove(drawing)

    # 删除文档关系中的图片引用（包括页眉页脚中的图片）
    for rel_id, rel in list(doc.part.rels.items()):
        if 'image' in rel.reltype:
            try:
                del doc.part.rels[rel_id]
            except KeyError:
                pass

    doc.save(doc_path)
    print(f"    已删除 {len(blip_elements)} 个图片")


def remove_header_footer(doc_path):
    """删除 docx 文件中的所有页眉页脚"""
    print(f"  - 删除页眉页脚：{doc_path}")

    doc = Document(doc_path)

    # 删除页眉页脚引用
    removed_count = 0
    for rel_id, rel in list(doc.part.rels.items()):
        if 'header' in rel.reltype or 'footer' in rel.reltype:
            try:
                del doc.part.rels[rel_id]
                removed_count += 1
            except KeyError:
                pass

    doc.save(doc_path)
    print(f"    已删除 {removed_count} 个页眉页脚引用")


def clear_first_page_contact_info(doc_path):
    """
    清除 0_报告首页及目录.docx 中的联系信息表格行
    删除：被测单位、单位名称、单位地址、联系人等信息
    """
    print(f"  - 清除联系信息：{doc_path}")

    doc = Document(doc_path)

    # 需要删除的关键词
    contact_keywords = ['被测单位', '单位名称', '单位地址', '联系人', '姓    名', '职务/职称',
                        '所属部门', '办公电话', '移动电话', '电子邮件', '邮政编码',
                        '海尔集团电子商务有限公司', '崂山区海尔路', '单小真', '18864805126']

    for table in doc.tables:
        rows_to_remove = []
        for i, row in enumerate(table.rows):
            row_text = ''.join([cell.text for cell in row.cells])
            if any(kw in row_text for kw in contact_keywords):
                rows_to_remove.append(i)

        # 从后向前删除，避免索引错乱
        for i in reversed(rows_to_remove):
            row = table.rows[i]
            tc = row._tr
            tc.getparent().remove(tc)

    doc.save(doc_path)
    print(f"    已清理联系信息行")


def replace_text_in_docx(doc_path, text_replacements, regex_replacements):
    """在 docx 文件中替换文本"""
    print(f"  - 文本替换：{doc_path}")

    doc = Document(doc_path)
    total_count = 0

    # 遍历所有段落
    for paragraph in doc.paragraphs:
        count = replace_in_paragraph(paragraph, text_replacements, regex_replacements)
        total_count += count

    # 遍历所有表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    count = replace_in_paragraph(paragraph, text_replacements, regex_replacements)
                    total_count += count

    doc.save(doc_path)
    print(f"    共替换 {total_count} 处")


def replace_in_paragraph(paragraph, text_replacements, regex_replacements):
    """在段落中替换文本，保持格式"""
    total_count = 0

    # 获取段落中的所有 runs
    runs = paragraph.runs

    if not runs:
        return 0

    # 收集所有 runs 的文本
    original_texts = [run.text for run in runs]
    full_text = ''.join(original_texts)

    # 执行文本替换
    new_text = full_text
    for old_val, new_val in text_replacements.items():
        count = new_text.count(old_val)
        if count > 0:
            new_text = new_text.replace(old_val, new_val)
            total_count += count

    # 执行正则替换
    for pattern, replacement in regex_replacements.items():
        matches = re.findall(pattern, new_text)
        count = len(matches)
        if count > 0:
            new_text = re.sub(pattern, replacement, new_text)
            total_count += count

    # 如果文本有变化，更新 runs
    if new_text != full_text:
        if len(runs) == 1:
            runs[0].text = new_text
        else:
            for i, run in enumerate(runs):
                if i < len(new_text):
                    chunk_size = max(1, len(new_text) // len(runs))
                    start = i * chunk_size
                    if i == len(runs) - 1:
                        run.text = new_text[start:]
                    else:
                        run.text = new_text[start:start + chunk_size]

    return total_count


def desensitize_documents(input_dir, recursive=True):
    """
    对指定目录下的所有 docx 文档进行脱敏处理

    Args:
        input_dir: 输入目录
        recursive: 是否递归处理子目录
    """
    input_dir = os.path.abspath(input_dir)

    if not os.path.exists(input_dir):
        print(f"错误：目录不存在 - {input_dir}")
        return

    # 定义文本替换规则（按长度降序排列，长文本优先替换）
    basic_replacements = {
        '山东新潮信息技术有限公司': '建筑工地',
        '山东省济南市高新区汉峪金谷 A3-4-16': '塞伯坦',
        '山东省济南市高新区汉峪金谷 A3-4-16': '塞伯坦',
        '济南市高新区汉峪金谷 A3-4-16': '塞伯坦',
        '济南市高新区汉峪金谷 A3-4-16': '塞伯坦',
        '青岛': '乌城',
        '山东': '乌城',
        '中国': '乌城',
        'SC202127130010134': '000000',
        '250102': '0000',
        '刘俊': '谪仙人',
        '测评服务部': '屠宰场',
        '0532-88981060': '000000',
        '18661855609': '0000000',
        'liujun@sdsecurity.org.cn': 'aaa',
        '0134': '9999',
    }

    # 定义正则替换规则
    regex_replacements = {
        r'\d{11}-\d{5}': '00-00',  # 匹配 37020036007-20001 这种格式
        r'济南市高新区汉峪金谷\s*A3-4-16': '塞伯坦',  # 匹配有空格和没空格的版本
    }

    # 查找所有 docx 文件
    docx_files = []
    if recursive:
        for root, dirs, files in os.walk(input_dir):
            for file in files:
                if file.endswith('.docx'):
                    docx_files.append(os.path.join(root, file))
    else:
        for file in os.listdir(input_dir):
            if file.endswith('.docx'):
                docx_files.append(os.path.join(input_dir, file))

    if not docx_files:
        print(f"未在目录中找到 docx 文件")
        return

    print(f"找到 {len(docx_files)} 个 docx 文件")
    print("=" * 50)

    for docx_file in docx_files:
        filename = os.path.basename(docx_file)
        print(f"\n处理文件：{filename}")
        print("-" * 40)

        # 1. 删除图片
        remove_images_from_docx(docx_file)

        # 2. 删除页眉页脚
        remove_header_footer(docx_file)

        # 3. 清除首页联系信息（仅首页文件）
        if filename == '0_报告首页及目录.docx':
            clear_first_page_contact_info(docx_file)

        # 4. 文本替换
        replace_text_in_docx(docx_file, basic_replacements, regex_replacements)

        print(f"完成：{filename}")

    print("\n" + "=" * 50)
    print(f"所有文件脱敏处理完成！")


def main():
    parser = argparse.ArgumentParser(
        description='文档脱敏工具 - 删除图片、页眉页脚，替换敏感信息',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python desensitize.py /path/to/docs          # 处理指定目录下的所有 docx 文件
  python desensitize.py /path/to/docs --no-recursive  # 仅处理目录，不递归子目录
        """
    )
    parser.add_argument(
        'input_dir',
        help='包含 docx 文件的目录路径'
    )
    parser.add_argument(
        '--no-recursive',
        action='store_true',
        help='不递归处理子目录'
    )

    args = parser.parse_args()

    desensitize_documents(
        args.input_dir,
        recursive=not args.no_recursive
    )


if __name__ == '__main__':
    main()
