#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
等级测评报告拆分脚本

将完整的等级测评报告 docx 文件按章节和附录拆分成多个独立文件，
便于大模型进行内容审核。

使用方法:
    python split_report.py [输入文件] [输出目录]

示例:
    python split_report.py report.docx ./chapters
"""

from docx import Document
from docx.oxml.ns import qn
import copy
import os
import sys
import zipfile
import shutil
import tempfile
import re
from pathlib import Path
from xml.etree import ElementTree as ET


# 章节定义：(起始段落，结束段落，章节名称)
# 如果为 None，则自动分析文档结构生成
CHAPTERS = None


def build_body_elements(doc):
    """构建文档 body 元素列表"""
    body_elements = []
    para_idx = 0

    for child in doc.element.body:
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            body_elements.append(('para', para_idx, child))
            para_idx += 1
        elif tag == 'tbl':
            body_elements.append(('table', -1, child))

    return body_elements


def collect_media_refs(element):
    """收集元素中引用的所有媒体 rId"""
    rIds = set()

    # 查找所有 blip 元素（无论命名空间）
    for blip in element.iter():
        if 'blip' in blip.tag and 'main' in blip.tag:
            rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if rId:
                rIds.add(rId)

    # 查找 v:imagedata 元素
    for img_data in element.iter():
        if 'imagedata' in img_data.tag:
            rId = img_data.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
            if rId:
                rIds.add(rId)

    return rIds


def split_chapter_with_media(src_doc_path, body_elements, start_para, end_para, chapter_name):
    """
    拆分章节并保留图片

    使用 zipfile 直接操作 docx 文件，复制相关的媒体文件
    """
    # 创建临时目录
    temp_dir = tempfile.mkdtemp()

    try:
        # 解压源文档
        with zipfile.ZipFile(src_doc_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)

        # 读取 document.xml
        doc_xml_path = os.path.join(temp_dir, 'word', 'document.xml')
        with open(doc_xml_path, 'r', encoding='utf-8') as f:
            doc_content = f.read()

        # 解析 XML
        root = ET.fromstring(doc_content)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        # 构建 body 元素列表
        src_elements = []
        para_idx = 0
        for child in root.findall('.//w:body/*', ns):
            tag = child.tag.split('}')[-1]
            if tag == 'p':
                src_elements.append(('para', para_idx, child))
                para_idx += 1
            elif tag == 'tbl':
                src_elements.append(('table', -1, child))

        # 创建新的 document.xml
        new_root = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}document')
        new_body = ET.SubElement(new_root, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}body')

        # 收集需要复制的媒体 rId（只收集章节内的）
        needed_rIds = set()
        current_para = 0

        for elem_type, p_idx, child in src_elements:
            if elem_type == 'para':
                if current_para >= start_para and current_para < end_para:
                    new_body.append(copy_xml_element(child))
                    needed_rIds.update(collect_media_refs(child))
                current_para += 1
            elif elem_type == 'table':
                if current_para > start_para and current_para <= end_para:
                    new_body.append(copy_xml_element(child))
                    needed_rIds.update(collect_media_refs(child))

        # 读取并更新关系文件
        rels_path = os.path.join(temp_dir, 'word', '_rels', 'document.xml.rels')
        new_rels_root = ET.Element('{http://schemas.openxmlformats.org/package/2006/relationships}Relationships')

        if os.path.exists(rels_path):
            with open(rels_path, 'r', encoding='utf-8') as f:
                rels_content = f.read()
            rels_root = ET.fromstring(rels_content)

            # 创建新的关系文件，只保留需要的引用
            for rel in rels_root.findall('{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
                rel_id = rel.get('Id')
                rel_type = rel.get('Type')

                # 保留图片关系（只保留章节内用到的）
                if 'image' in rel_type:
                    if rel_id in needed_rIds:
                        new_rels_root.append(copy_xml_element(rel))
                # 保留其他必要的关系（如 header, footer, styles 等）
                else:
                    new_rels_root.append(copy_xml_element(rel))

        # 写入新的关系文件
        with open(rels_path, 'w', encoding='utf-8') as f:
            f.write(ET.tostring(new_rels_root, encoding='unicode'))

        # 写入新的 document.xml
        with open(doc_xml_path, 'w', encoding='utf-8') as f:
            f.write(ET.tostring(new_root, encoding='unicode'))

        # 删除不需要的图片文件以减小体积
        media_dir = os.path.join(temp_dir, 'word', 'media')
        if os.path.exists(media_dir):
            # 获取关系文件中引用的图片文件名
            needed_files = set()
            for rel in new_rels_root.findall('{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
                rel_type = rel.get('Type')
                target = rel.get('Target')
                if 'image' in rel_type and target:
                    # Target 类似 media/image1.png
                    needed_files.add(target.split('/')[-1])

            # 删除不需要的图片
            for file in os.listdir(media_dir):
                if file not in needed_files:
                    os.remove(os.path.join(media_dir, file))

        # 创建输出文件
        output_file = os.path.join(temp_dir, 'output.docx')
        with zipfile.ZipFile(output_file, 'w', zipfile.ZIP_DEFLATED) as zip_out:
            for root, dirs, files in os.walk(temp_dir):
                for file in files:
                    if file == 'output.docx':
                        continue
                    file_path = os.path.join(root, file)
                    arc_name = os.path.relpath(file_path, temp_dir)
                    zip_out.write(file_path, arc_name)

        # 移动到目标位置
        return output_file, temp_dir

    except Exception as e:
        shutil.rmtree(temp_dir, ignore_errors=True)
        raise e


def copy_xml_element(element):
    """深拷贝 XML 元素"""
    return ET.fromstring(ET.tostring(element))


def create_header_footer_doc(doc):
    """
    创建包含所有页眉页脚内容的文档

    Args:
        doc: 原始文档对象

    Returns:
        Document: 页眉页脚文档
    """
    hf_doc = Document()
    hf_doc.add_heading('文档页眉页脚内容', 0)

    for i, section in enumerate(doc.sections):
        hf_doc.add_heading(f'节 {i+1}', level=1)

        hf_doc.add_heading('页眉:', level=2)
        try:
            for p in section.header.paragraphs:
                if p.text.strip():
                    hf_doc.add_paragraph(p.text.strip())
        except Exception:
            hf_doc.add_paragraph("无法读取")

        hf_doc.add_heading('页脚:', level=2)
        try:
            for p in section.footer.paragraphs:
                if p.text.strip():
                    hf_doc.add_paragraph(p.text.strip())
        except Exception:
            hf_doc.add_paragraph("无法读取")

    return hf_doc


def count_content(doc):
    """
    统计文档内容（包括表格）

    Args:
        doc: Document 对象

    Returns:
        tuple: (字符数，估算 token 数)
    """
    chars = 0

    # 段落
    for para in doc.paragraphs:
        chars += len(para.text.strip())

    # 表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chars += len(cell.text.strip())

    return chars, int(chars * 1.5)


def auto_detect_chapters(doc):
    """
    自动分析文档结构，检测章节位置

    策略：
    1. 只使用实际内容标题（Heading 1 和附录一级标题）
    2. 忽略 toc 1 目录索引标题
    3. 附录内容按附录字母（A、B、C...）合并，不拆分表格

    Args:
        doc: Document 对象

    Returns:
        list: 章节列表 [(start_para, end_para, chapter_name), ...]
    """
    chapters = []
    heading_positions = []
    content_headings = []  # 实际正文标题 (Heading 1)
    appendix_headings = []  # 附录主标题（附录一级标题）

    # 1. 遍历所有段落，只收集实际内容标题
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ''
        text = para.text.strip()

        if not text:
            continue

        # 只收集 Heading 1（主报告章节）
        if 'Heading 1' in style_name:
            content_headings.append((i, text, style_name))
            heading_positions.append((i, text, style_name))
        # 收集附录一级标题（实际附录内容开始）
        elif '附录一级标题' in style_name:
            appendix_headings.append((i, text, style_name))
            heading_positions.append((i, text, '附录'))

    # 2. 构建章节列表
    # 2a. 添加封面和目录（第一个 Heading 1 之前的内容）
    if content_headings:
        first_heading_pos = content_headings[0][0]
        if first_heading_pos > 0:
            chapters.append((0, first_heading_pos, '0_报告首页及目录'))

    # 2b. 添加主报告章节
    chapter_num = 1
    for i, (pos, text, style) in enumerate(content_headings):
        # 确定结束位置
        if i + 1 < len(content_headings):
            end_pos = content_headings[i + 1][0]
        else:
            # 最后一个主报告章节，结束位置为附录开始
            if appendix_headings:
                end_pos = appendix_headings[0][0]
            else:
                end_pos = len(doc.paragraphs)

        # 清理章节名称
        chapter_name = f"{chapter_num}_{text[:30]}"
        chapters.append((pos, end_pos, chapter_name))
        chapter_num += 1

    # 2c. 添加附录章节（按附录一级标题顺序）
    for i, (pos, text, style) in enumerate(appendix_headings):
        if i + 1 < len(appendix_headings):
            end_pos = appendix_headings[i + 1][0]
        else:
            end_pos = len(doc.paragraphs)

        # 生成附录名称 - 清理特殊字符
        safe_name = re.sub(r'[\\/:*?"<>|]', '_', text)
        appendix_name = f"附录_{safe_name[:40]}"
        chapters.append((pos, end_pos, appendix_name))

    return chapters, heading_positions


def clean_appendix_name(text):
    """
    清理附录名称，生成合适的文件名

    Args:
        text: 附录标题文本

    Returns:
        str: 清理后的名称
    """
    # 移除特殊字符和空格
    name = re.sub(r'[\\/:*?"<>|]', '_', text)
    name = re.sub(r'\s+', '_', name)
    # 限制长度
    if len(name) > 50:
        name = name[:50]
    return name


def analyze_document(input_file):
    """
    分析文档结构，输出章节信息

    Args:
        input_file: 输入文件路径
    """
    print("=" * 60)
    print("文档结构分析")
    print("=" * 60)

    doc = Document(input_file)

    # 统计总体信息
    total_chars, total_tokens = count_content(doc)
    print(f"\n原文档统计:")
    print(f"  段落数：{len(doc.paragraphs)}")
    print(f"  表格数：{len(doc.tables)}")
    print(f"  节数：{len(doc.sections)}")
    print(f"  总字符数：{total_chars:,}")
    print(f"  估算 token 数：~{total_tokens:,}")

    # 自动检测章节
    print("\n自动检测的章节结构:")
    chapters, heading_positions = auto_detect_chapters(doc)

    for i, (start, end, name) in enumerate(chapters):
        print(f"  {i}. {name}: 段落 {start}-{end} ({end - start} 段)")

    # 显示所有检测到的标题
    print("\n检测到的标题位置:")
    for pos, text, style in heading_positions:
        preview = text[:50] + "..." if len(text) > 50 else text
        print(f"  段落 {pos} [{style}]: {preview}")

    return doc, chapters


def split_document(input_file, output_dir, chapters_config=None):
    """
    拆分文档为主函数

    Args:
        input_file: 输入文件路径
        output_dir: 输出目录
        chapters_config: 章节配置，如果为 None 则自动检测
    """
    # 确保输出目录存在
    os.makedirs(output_dir, exist_ok=True)

    print("=" * 60)
    print("开始拆分文档")
    print("=" * 60)
    print(f"\n输入文件：{input_file}")
    print(f"输出目录：{output_dir}\n")

    # 加载文档
    doc = Document(input_file)

    # 自动检测或手动指定章节配置
    if chapters_config is None:
        print("正在自动分析文档结构...")
        _, auto_chapters = auto_detect_chapters(doc)
        chapters_config = auto_chapters
        print(f"检测到 {len(chapters_config)} 个章节\n")
    else:
        print(f"使用手动配置的 {len(chapters_config)} 个章节\n")

    # 构建 body 元素列表（用于确定章节范围）
    body_elements = build_body_elements(doc)
    print(f"文档元素：{len(body_elements)} 个（{len(doc.paragraphs)} 段落 + {len(doc.tables)} 表格）")

    # 拆分章节
    print(f"\n正在拆分章节...")
    files_created = []
    temp_dirs = []  # 记录临时目录

    for start_para, end_para, chapter_name in chapters_config:
        try:
            # 使用 zipfile 方法拆分章节（包含图片）
            temp_file, temp_dir = split_chapter_with_media(input_file, body_elements, start_para, end_para, chapter_name)
            temp_dirs.append(temp_dir)

            # 移动到目标位置
            output_file = os.path.join(output_dir, f'{chapter_name}.docx')
            shutil.move(temp_file, output_file)

            # 统计信息
            chapter_doc = Document(output_file)
            chars, tokens = count_content(chapter_doc)
            files_created.append((chapter_name, tokens))

            print(f"  ✓ {chapter_name}.docx ({end_para - start_para} 段落，~{tokens:,} tokens)")
        except Exception as e:
            print(f"  ✗ {chapter_name}.docx 失败：{e}")

    # 清理临时目录
    for temp_dir in temp_dirs:
        shutil.rmtree(temp_dir, ignore_errors=True)

    # 创建页眉页脚文档
    print("\n正在生成页眉页脚文档...")
    hf_doc = create_header_footer_doc(doc)
    hf_file = os.path.join(output_dir, '页眉页脚.docx')
    hf_doc.save(hf_file)
    hf_chars, hf_tokens = count_content(hf_doc)
    files_created.append(('页眉页脚', hf_tokens))
    print(f"  ✓ 页眉页脚.docx (~{hf_tokens:,} tokens)")

    # 输出统计
    print("\n" + "=" * 60)
    print("拆分完成!")
    print("=" * 60)

    total_tokens = sum(t for _, t in files_created)
    print(f"\n共生成 {len(files_created)} 个文件:")

    # 按 token 数排序
    sorted_files = sorted(files_created, key=lambda x: x[1], reverse=True)
    for name, tokens in sorted_files:
        status = "✓" if tokens < 200000 else "⚠ 接近限制"
        print(f"  {status} {name}.docx: ~{tokens:,} tokens")

    print(f"\n累计总 token: ~{total_tokens:,}")
    print(f"Claude 200K 限制：{'✓ 可分批处理' if total_tokens < 200000 * 2 else '⚠ 需要更多拆分'}")
    print(f"\n文件保存位置：{os.path.abspath(output_dir)}")


def main():
    # 默认路径
    default_input = 'report.docx'
    default_output = './chapters'

    # 解析命令行参数
    import argparse
    parser = argparse.ArgumentParser(description="等级测评报告拆分脚本")
    parser.add_argument("input_file", nargs="?", default=None, help="输入的 docx 文件路径")
    parser.add_argument("output_dir", nargs="?", default=None, help="输出目录")
    parser.add_argument("--auto", action="store_true", help="使用自动检测的章节配置，不询问")
    parser.add_argument("--use-chapters", action="store_true", help="使用脚本中硬编码的 CHAPTERS 配置")
    parser.add_argument("--skip-desensitize", action="store_true", help="跳过脱敏处理步骤")

    args = parser.parse_args()

    input_file = args.input_file if args.input_file else default_input
    output_dir = args.output_dir if args.output_dir else default_output

    # 检查输入文件
    if not os.path.exists(input_file):
        print(f"错误：找不到输入文件 '{input_file}'")
        print(f"\n使用方法:")
        print(f"  python {sys.argv[0]} [输入文件] [输出目录]")
        print(f"\n示例:")
        print(f"  python {sys.argv[0]} report.docx ./chapters")
        sys.exit(1)

    # 先分析文档并获取自动检测的章节
    doc, auto_chapters = analyze_document(input_file)
    print("\n")

    # 根据参数决定使用哪个配置
    if args.auto or args.use_chapters:
        # 非交互模式
        if args.use_chapters and CHAPTERS is not None:
            print("使用脚本中硬编码的 CHAPTERS 配置")
            chapters_config = CHAPTERS
        else:
            print("使用自动检测的章节配置")
            chapters_config = auto_chapters
    else:
        # 交互模式
        print("是否使用自动检测的章节配置进行拆分？")
        print("  Y - 使用自动检测的配置")
        print("  N - 使用脚本中硬编码的 CHAPTERS 配置")
        print("  Q - 退出不执行")
        print()

        choice = input("请选择 [Y/N/Q]: ").strip().upper()

        if choice == 'Q':
            print("已取消操作")
            sys.exit(0)
        elif choice == 'N':
            if CHAPTERS is None:
                print("警告：脚本中没有配置 CHAPTERS，将使用自动检测的配置")
                chapters_config = auto_chapters
            else:
                print("使用脚本中硬编码的 CHAPTERS 配置")
                chapters_config = CHAPTERS
        else:
            print("使用自动检测的章节配置")
            chapters_config = auto_chapters

    print("\n")

    # 拆分文档
    split_document(input_file, output_dir, chapters_config)

    # 复制高风险判定指引文件到输出目录
    script_dir = os.path.dirname(os.path.abspath(__file__))
    guidance_file = os.path.join(script_dir, '高风险判定指引.md')
    if os.path.exists(guidance_file):
        import shutil
        target_file = os.path.join(output_dir, '高风险判定指引.md')
        shutil.copy2(guidance_file, target_file)
        print(f"已复制高风险判定指引.md 到 {output_dir}")
    else:
        print(f"警告：找不到高风险判定指引.md - {guidance_file}")

    # 复制 beian.png 到输出目录
    beian_file = os.path.join(script_dir, 'beian.png')
    if os.path.exists(beian_file):
        target_png = os.path.join(output_dir, 'beian.png')
        shutil.copy2(beian_file, target_png)
        print(f"已复制 beian.png 到 {output_dir}")
    else:
        print(f"警告：找不到 beian.png - {beian_file}")

    # 自动调用脱敏脚本（如果未指定跳过）
    if not args.skip_desensitize:
        print("\n正在对拆分后的文件进行脱敏处理...")
        # desensitize.py 在 notebooklm_auto_review 目录中
        desensitize_script = os.path.join(script_dir, 'notebooklm_auto_review', 'desensitize.py')

        if os.path.exists(desensitize_script):
            import subprocess
            try:
                result = subprocess.run(
                    ['python3', desensitize_script, output_dir],
                    capture_output=True,
                    text=True,
                    timeout=300
                )
                print(result.stdout)
                if result.stderr:
                    print(f"警告：{result.stderr}")
                print("脱敏处理完成！")
            except subprocess.TimeoutExpired:
                print("错误：脱敏处理超时")
            except Exception as e:
                print(f"错误：脱敏处理失败 - {e}")
        else:
            print(f"警告：找不到脱敏脚本 - {desensitize_script}")
    else:
        print("\n已跳过脱敏处理步骤")


if __name__ == '__main__':
    main()
